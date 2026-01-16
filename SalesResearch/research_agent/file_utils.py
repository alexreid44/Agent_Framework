"""File generation utilities for PDF and Word documents."""

from pathlib import Path
from datetime import datetime
import re
from html.parser import HTMLParser


def generate_filename(company_name: str, suffix: str = "") -> tuple[str, str]:
    """Generate timestamped filename and sanitized company name.

    Args:
        company_name: The company name to sanitize
        suffix: Optional suffix to add (e.g., "_BusinessModelCanvas")

    Returns:
        Tuple of (base_filename, timestamp)
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_company_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in company_name)
    safe_company_name = safe_company_name.replace(' ', '_')
    base_filename = f"{timestamp}_{safe_company_name}{suffix}"
    return base_filename, timestamp


def save_research_pdf(output_dir: Path, base_filename: str, company_name: str,
                     company_url: str, content: str) -> Path:
    """Save research report as PDF using ReportLab.

    Args:
        output_dir: Directory to save the file
        base_filename: Base filename without extension
        company_name: Company name for title
        company_url: Company URL for metadata
        content: Research content in HTML format

    Returns:
        Path to the saved PDF file
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from html import unescape

    pdf_file = output_dir / f"{base_filename}.pdf"

    # Create PDF document
    doc = SimpleDocTemplate(
        str(pdf_file),
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=1*inch,
        bottomMargin=0.75*inch
    )

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#2c3e50',
        spaceAfter=30,
        alignment=TA_CENTER
    )

    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#2c3e50',
        spaceAfter=12,
        spaceBefore=12
    )

    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='#34495e',
        spaceAfter=10,
        spaceBefore=10
    )

    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor='#7f8c8d',
        spaceAfter=8,
        spaceBefore=8
    )

    # Add title
    elements.append(Paragraph(f"Strategic Pursuit Plan: {company_name}", title_style))
    elements.append(Spacer(1, 0.2*inch))

    # Add metadata
    metadata_style = ParagraphStyle(
        'Metadata',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#7f8c8d'
    )
    elements.append(Paragraph(f"<b>Company URL:</b> {company_url}", metadata_style))
    elements.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", metadata_style))
    elements.append(Spacer(1, 0.3*inch))

    # Parse HTML content and convert to ReportLab elements
    # Clean up the HTML and convert to ReportLab-compatible format
    content = unescape(content)

    # Split by common HTML tags and process
    import re
    # Simple regex-based HTML parsing for ReportLab
    html_parts = re.split(r'(<h[1-6]>.*?</h[1-6]>|<p>.*?</p>|<ul>.*?</ul>|<ol>.*?</ol>|<li>.*?</li>)', content, flags=re.DOTALL)

    for part in html_parts:
        part = part.strip()
        if not part:
            continue

        # Headers
        if re.match(r'<h2>', part):
            text = re.sub(r'<.*?>', '', part)
            if text.strip():
                elements.append(Paragraph(text, heading1_style))
        elif re.match(r'<h3>', part):
            text = re.sub(r'<.*?>', '', part)
            if text.strip():
                elements.append(Paragraph(text, heading2_style))
        elif re.match(r'<h4>', part):
            text = re.sub(r'<.*?>', '', part)
            if text.strip():
                elements.append(Paragraph(text, heading3_style))
        # Paragraphs
        elif re.match(r'<p>', part):
            # Keep basic HTML tags for ReportLab
            text = part.replace('<p>', '').replace('</p>', '').strip()
            if text:
                elements.append(Paragraph(text, styles['Normal']))
                elements.append(Spacer(1, 0.1*inch))
        # List items
        elif re.match(r'<li>', part):
            text = part.replace('<li>', '').replace('</li>', '').strip()
            if text:
                elements.append(Paragraph(f"• {text}", styles['Normal']))
        # Plain text fallback
        elif part and not re.match(r'<[^>]+>', part):
            elements.append(Paragraph(part, styles['Normal']))

    # Build PDF
    doc.build(elements)
    return pdf_file


class HTMLToDocxParser(HTMLParser):
    """Parse HTML and convert to Word document structure."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.in_list = False
        self.in_bold = False
        self.in_italic = False
        self.in_heading = False
        self.heading_level = 0
        self.heading_text = ""

    def handle_starttag(self, tag, attrs):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_heading = True
            self.heading_level = int(tag[1])
            self.heading_text = ""
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
        elif tag in ['ul', 'ol']:
            self.in_list = True
        elif tag == 'li':
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
        elif tag == 'strong' or tag == 'b':
            self.in_bold = True
        elif tag == 'em' or tag == 'i':
            self.in_italic = True
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')

    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if self.heading_text.strip():
                self.doc.add_heading(self.heading_text.strip(), level=self.heading_level)
            self.in_heading = False
            self.heading_text = ""
        elif tag == 'p':
            self.current_paragraph = None
        elif tag in ['ul', 'ol']:
            self.in_list = False
        elif tag == 'li':
            self.current_paragraph = None
        elif tag == 'strong' or tag == 'b':
            self.in_bold = False
        elif tag == 'em' or tag == 'i':
            self.in_italic = False

    def handle_data(self, data):
        if not data.strip():
            return

        if self.in_heading:
            self.heading_text += data
        elif self.current_paragraph:
            run = self.current_paragraph.add_run(data)
            if self.in_bold:
                run.bold = True
            if self.in_italic:
                run.italic = True


def save_business_model_canvas_docx(output_dir: Path, base_filename: str,
                                   company_name: str, company_url: str,
                                   content: str) -> Path:
    """Save Business Model Canvas as Word document from HTML content.

    Args:
        output_dir: Directory to save the file
        base_filename: Base filename without extension
        company_name: Company name for title
        company_url: Company URL for metadata
        content: Business Model Canvas HTML content

    Returns:
        Path to the saved DOCX file
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx_file = output_dir / f"{base_filename}.docx"
    doc = Document()

    # Title
    title = doc.add_heading(f"Business Model Canvas: {company_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Company URL: ").bold = True
    metadata.add_run(company_url)

    timestamp_para = doc.add_paragraph()
    timestamp_para.add_run(f"Generated: ").bold = True
    timestamp_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    doc.add_paragraph()  # Empty line
    doc.add_paragraph("_" * 80)  # Separator
    doc.add_paragraph()

    # Parse HTML content and add to document
    parser = HTMLToDocxParser(doc)
    parser.feed(content)

    doc.save(str(docx_file))
    return docx_file
